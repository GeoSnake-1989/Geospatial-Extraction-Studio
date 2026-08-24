from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).with_name("Geospatial-Extraction-Studio-User-Manual.docx")
BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 99, 110)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_font(run, size=11, bold=None, color=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if italic is not None:
        run.italic = italic


def add_body(doc, text, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    set_font(p.add_run(text), 11, italic=italic)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    set_font(p.add_run(text), 11)
    return p


def add_step(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    set_font(p.add_run(text), 11)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    fmt = p.paragraph_format
    fmt.space_before = Pt(18 if level == 1 else 12)
    fmt.space_after = Pt(8 if level == 1 else 6)
    run = p.add_run(text)
    set_font(run, 16 if level == 1 else 13, bold=True, color=BLUE if level == 1 else DARK)
    return p


def add_note(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    cell = table.cell(0, 0)
    shade(cell, "F4F6F9")
    set_cell_margins(cell, 100, 160, 100, 160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    lead = p.add_run(title + ": ")
    set_font(lead, 10.5, bold=True, color=DARK)
    set_font(p.add_run(text), 10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    table.autofit = False
    for i, (header, width) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[i]
        cell.width = Inches(width)
        shade(cell, "E8EEF5")
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(header), 10, bold=True, color=DARK)
    for row in rows:
        cells = table.add_row().cells
        for i, (value, width) in enumerate(zip(row, widths)):
            cells[i].width = Inches(width)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[i])
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_font(p.add_run(value), 9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def main():
    doc = Document()
    doc.core_properties.author = "Jacob Horwitz"
    doc.core_properties.title = "Geospatial Extraction Studio User Manual"
    doc.core_properties.subject = "User manual for Geospatial Extraction Studio"
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    styles["Normal"].font.size = Pt(11)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run("Geospatial Extraction Studio | User Manual"), 8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("Copyright 2026 Jacob Horwitz | Apache-2.0"), 8.5, color=MUTED)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(46)
    title.paragraph_format.space_after = Pt(10)
    set_font(title.add_run("Geospatial Extraction Studio"), 28, bold=True, color=DARK)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(22)
    set_font(subtitle.add_run("User Manual"), 17, color=BLUE)
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro.paragraph_format.space_after = Pt(24)
    set_font(intro.add_run("A practical guide to creating local terrain, NAIP aerial-imagery, and OpenStreetMap GIS outputs."), 11, color=MUTED)
    legal = doc.add_paragraph()
    legal.alignment = WD_ALIGN_PARAGRAPH.CENTER
    legal.paragraph_format.space_after = Pt(12)
    set_font(legal.add_run("Copyright 2026 Jacob Horwitz. Licensed under the Apache License, Version 2.0. Third-party data and software remain subject to their own notices and terms."), 8.5, color=MUTED)
    add_note(doc, "What this app does", "Use one place search and shared map to define an area, then choose the elevation, aerial imagery, or OpenStreetMap workflow. Outputs and their source evidence are saved locally.")
    doc.add_page_break()

    add_heading(doc, "1. Getting started")
    add_body(doc, "Geospatial Extraction Studio is a Windows local-first application. It can be run from source with Python and Node.js, or from the packaged installer. It does not require an OpenAI connection.")
    add_heading(doc, "Start the application", 2)
    add_step(doc, "From the project folder, double-click start.bat (or run start.ps1 in PowerShell).")
    add_step(doc, "Your browser opens the local application. If it does not, open the local address shown by the launcher (normally http://127.0.0.1:5173).")
    add_step(doc, "When finished, close the browser and use stop.bat or stop.ps1 to stop the local services.")
    add_note(doc, "Internet access", "The app itself runs locally, but place search and data extraction contact the selected public providers. Availability, rate limits, and response times are provider-controlled.")
    add_heading(doc, "The workspace at a glance", 2)
    add_table(doc, ["Area", "Purpose"], [
        ("Top bar", "Open Legal notices, view local Storage, and see the current data provider."),
        ("Workflow tabs", "Choose Elevation, Aerial imagery, or OpenStreetMap."),
        ("Left panel", "Search, define the extraction area, choose settings, and start an extraction."),
        ("Map", "Shows the selected area. Use the rectangle button to draw a new area."),
        ("Results panel", "Shows the 3D terrain viewer, aerial preview, or OSM result summary and downloads."),
    ], [1.55, 4.95])

    add_heading(doc, "2. Choose a location and area")
    add_body(doc, "All three workflows use the same selection controls, but their search coverage differs: elevation and NAIP are U.S.-focused, while OpenStreetMap search works worldwide.")
    add_heading(doc, "Search for a place", 2)
    add_step(doc, "Enter a city, landmark, or address in Find a location and select Go.")
    add_step(doc, "Choose a result. The map moves to it and the coordinate fields update.")
    add_step(doc, "For OSM, a named Polygon or MultiPolygon place may be retained as the selection boundary. For the other workflows, use the displayed rectangular extent.")
    add_heading(doc, "Draw or edit a rectangle", 2)
    add_step(doc, "Select the rectangle tool on the map.")
    add_step(doc, "Drag diagonally to draw the area. Press Esc or select the tool again to cancel.")
    add_step(doc, "Fine-tune west, north, east, and south directly in the coordinate fields if needed.")
    add_note(doc, "Selection limits", "Elevation and manually drawn OSM rectangles are limited to 2,500 km². A named OSM boundary can be up to 10,000 km². NAIP availability searches are limited to 500 km² and output size is also limited by native-resolution pixels.")

    add_heading(doc, "3. Extract elevation and explore terrain")
    add_body(doc, "The Elevation workflow retrieves a bounded DEM from the USGS 3DEP seamless service, preserves the source GeoTIFF, and creates a lightweight local terrain preview.")
    add_step(doc, "Select the Elevation tab, search for a U.S. location or draw an area, and set a dataset label.")
    add_step(doc, "Choose 3D preview detail: Compact (64 × 64), Balanced (96 × 96), or Detailed (128 × 128). This only changes viewer performance; it does not change the source download grid.")
    add_step(doc, "Select Extract elevation and wait for the queued, downloading, and processing status to complete.")
    add_step(doc, "Use the 3D viewer to inspect terrain. Adjust vertical exaggeration from 0× (flat) to 4×; 1× preserves real-world vertical-to-horizontal scale.")
    add_step(doc, "Optionally enable Mesh, toggle the 0 m Sea level reference plane, or choose Topo, Ember, or Glacial coloring.")
    add_body(doc, "The result summary reports low, mean, high, and relief in meters. It also exposes CRS, source and preview dimensions, and vertical datum. Any unavailable CRS, NoData, unit, or vertical-datum value remains explicitly not declared.")
    add_heading(doc, "Download elevation results", 2)
    add_bullet(doc, "Download DEM + sources: a ZIP containing the DEM GeoTIFF, source evidence JSON, and readable source documentation.")
    add_bullet(doc, "Source GeoTIFF: the original downloaded elevation raster.")
    add_bullet(doc, "Source documentation: readable provider and processing evidence.")

    add_heading(doc, "4. Extract NAIP aerial imagery")
    add_body(doc, "The Aerial imagery workflow searches the configured USDA NAIP catalog through the Planetary Computer STAC service. It writes a clipped AOI GeoTIFF, PNG preview, source-item manifest, and source documentation.")
    add_step(doc, "Select Aerial imagery and define a U.S. rectangular area. Availability is checked automatically.")
    add_step(doc, "Choose Imagery date: Latest complete coverage for a single, consistent acquisition year; a historical complete year; or Newest per tile to fill each part with the newest published tile.")
    add_step(doc, "Choose Natural color (RGB) or Four band (RGB + near infrared).")
    add_step(doc, "Set the label and select Extract NAIP imagery. The button is disabled while coverage is being checked or if the estimated native-resolution image is too large.")
    add_body(doc, "After completion, review the preview, acquisition-date range, pixel dimensions, number of bands, resolution, CRS, coverage, attribution, and license record.")
    add_note(doc, "Date seams", "Latest complete coverage uses one year with at least 99.5% coverage. Newest per tile can mix acquisition dates and can show visible seams.")
    add_heading(doc, "Download imagery results", 2)
    add_bullet(doc, "Download imagery + sources: ZIP package with imagery and source records.")
    add_bullet(doc, "Download AOI GeoTIFF: the clipped, tiled and compressed output raster.")
    add_bullet(doc, "Source manifest and Source documentation: catalog items, source assets, dates, resolution, CRS, licenses, and selected tiles.")

    add_heading(doc, "5. Export OpenStreetMap features")
    add_body(doc, "The OpenStreetMap workflow queries OSM through OSMnx/Overpass, separates mixed geometry into GIS-ready feature classes, and supplies mandatory ODbL attribution with every archive.")
    add_step(doc, "Select OpenStreetMap and search worldwide, draw a rectangle, or select a named place boundary.")
    add_step(doc, "Choose an output format: OpenFileGDB or GeoPackage.")
    add_step(doc, "Create a new dataset, or choose a compatible recent dataset to add a later tag extraction while preserving existing layers.")
    add_step(doc, "Set OSM feature type (the tag key) and, if useful, an optional subtype/tag value. For example, use building with a blank subtype for all buildings, or amenity with school for schools.")
    add_step(doc, "Select Build [format], or Add feature classes when appending to an existing dataset.")
    add_body(doc, "When completed, the map can display the extracted features and show pop-ups with available attributes. The result card reports feature count, output classes, area, format, and license.")
    add_note(doc, "Layer names", "The app rejects an addition that would create a layer name already present in the target dataset. Choose another subtype or create a new output dataset.")
    add_heading(doc, "Download OSM results", 2)
    add_body(doc, "Select Download [format] + attribution to receive one ZIP containing the dataset and its OSM attribution/ODbL notice. Keep the notice with any redistributed OSM-derived data.")

    add_heading(doc, "6. Reopen, manage, and remove local outputs")
    add_heading(doc, "Recent outputs", 2)
    add_body(doc, "The Recent terrain, Recent NAIP, and Recent OSM buttons reopen up to three recent results for the active workflow. Completed history persists locally even though an in-progress job status resets when the app restarts.")
    add_heading(doc, "Storage", 2)
    add_step(doc, "Select Storage in the top bar to see total local storage and the terrain, imagery, OSM-export, and cache portions.")
    add_step(doc, "Use Clear cache to remove reusable OSM request-cache data. Completed terrain, NAIP imagery, and OSM datasets are preserved.")
    add_step(doc, "Use a trash button beside a saved item to permanently delete that saved output and remove it from local history. Confirm only when you no longer need the files.")
    add_note(doc, "Local locations", "Terrain source files and evidence are saved under data/original/terrain; processed terrain previews under data/processed; NAIP materials under data/processed/naip; OSM datasets and packages under data/exports. The SQLite history is data/app.db.")

    add_heading(doc, "7. Metadata, limits, and troubleshooting")
    add_heading(doc, "Metadata and source evidence", 2)
    add_body(doc, "Treat every output's source documentation and manifest/evidence files as part of the deliverable. Coordinate reference system, NoData, horizontal units, vertical units, and vertical datum are reported only when provided; the app does not infer missing metadata. The Sea level control is a visual 0 m reference, not a vertical-datum conversion.")
    add_heading(doc, "Known limits", 2)
    add_bullet(doc, "Elevation uses USGS 3DEP seamless only and is U.S.-focused.")
    add_bullet(doc, "NAIP is U.S.-focused; catalog publication, service availability, and pixel limits constrain requests.")
    add_bullet(doc, "Public Overpass queries can time out for dense or relation-heavy selections. Narrow the area or specify a tag value and try again.")
    add_bullet(doc, "The app currently does not support polygon clipping, user-supplied LiDAR, point-cloud viewing, general vertical-datum transformation, or multi-user deployment.")
    add_heading(doc, "If an extraction fails", 2)
    add_bullet(doc, "Read the displayed job message first; it distinguishes queued, downloading, processing, completed, and failed work.")
    add_bullet(doc, "Reduce the selection area, especially for NAIP or dense OSM categories.")
    add_bullet(doc, "For NAIP, choose a smaller rectangle or a different complete year if the native pixel estimate exceeds the configured limit.")
    add_bullet(doc, "For OSM, use a more specific subtype, such as amenity=school, and retry later if the public provider is busy.")
    add_bullet(doc, "Verify that the local backend is running and that the computer can reach the required data providers.")

    add_heading(doc, "Quick workflow checklist")
    add_table(doc, ["Goal", "Recommended path"], [
        ("Explore terrain", "Elevation → U.S. area → choose preview detail → Extract elevation → inspect 3D viewer → Download DEM + sources."),
        ("Obtain current aerial coverage", "Aerial imagery → U.S. rectangle → Latest complete coverage → choose bands → Extract → Download imagery + sources."),
        ("Build a GIS layer", "OpenStreetMap → worldwide area or boundary → tag key/value → GeoPackage or OpenFileGDB → Build → Download with attribution."),
        ("Free disk space safely", "Storage → Clear cache removes only reusable OSM cache; use a specific trash button only for outputs you intend to permanently remove."),
    ], [1.75, 4.75])

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
